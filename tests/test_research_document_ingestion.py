from __future__ import annotations

import sqlite3
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from nika_core.data.schema import MIGRATIONS, SCHEMA_VERSION
from nika_core.data.sqlite import SQLiteStore
from nika_core.research import (
    BlobStoreError,
    ChunkPolicy,
    ContentAddressedBlobStore,
    DocumentLimits,
    DocumentSecurityError,
    DocumentTooLargeError,
    ExtractionStatus,
    LocalCorpusService,
    ResearchRepository,
    ResearchWorkspace,
    SourceKind,
    SourceSpec,
    chunk_text,
    extract_document_file,
    preflight_office_archive,
)


def _repo(tmp_path: Path) -> tuple[SQLiteStore, ResearchRepository, Path]:
    store = SQLiteStore(tmp_path / "nika.db")
    store.initialize()
    repository = ResearchRepository(store)
    repository.upsert_workspace(ResearchWorkspace("ws", "Research"))
    source_root = tmp_path / "sources"
    source_root.mkdir()
    return store, repository, source_root


def _minimal_text_pdf(text: str = "Hello PDF") -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


def test_migration_10_applies_after_real_schema_9(tmp_path: Path) -> None:
    path = tmp_path / "v9.db"
    conn = sqlite3.connect(path)
    conn.execute(
        "CREATE TABLE schema_migrations(version INTEGER PRIMARY KEY, applied_at TEXT NOT NULL)"
    )
    for version in range(1, 10):
        for statement in MIGRATIONS[version]:
            conn.execute(statement)
        conn.execute(
            "INSERT INTO schema_migrations(version, applied_at) VALUES (?, 'fixture')",
            (version,),
        )
    conn.commit()
    conn.close()

    store = SQLiteStore(path)
    store.initialize()

    assert 10 in MIGRATIONS
    assert store.schema_version() == SCHEMA_VERSION
    with store.connection() as check:
        names = {
            row["name"]
            for row in check.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name LIKE 'corpus_%'"
            )
        }
    assert {
        "corpus_artifacts",
        "corpus_artifact_origins",
        "corpus_extractions",
        "corpus_document_artifacts",
    } <= names


def test_pdf_text_and_ocr_needed_classification(tmp_path: Path) -> None:
    text_pdf = tmp_path / "text.pdf"
    text_pdf.write_bytes(_minimal_text_pdf())
    extracted = extract_document_file(text_pdf)
    assert extracted.status is ExtractionStatus.EXTRACTED
    assert "Hello PDF" in extracted.text

    blank_pdf = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with blank_pdf.open("wb") as handle:
        writer.write(handle)
    blank = extract_document_file(blank_pdf)
    assert blank.status is ExtractionStatus.OCR_NEEDED
    assert blank.text == ""


def test_docx_and_xlsx_extract_structured_text(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Український документ")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Назва"
    table.cell(0, 1).text = "Грант"
    document.save(docx_path)
    docx_result = extract_document_file(docx_path)
    assert "Український документ" in docx_result.text
    assert "Назва\tГрант" in docx_result.text

    xlsx_path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Дані"
    sheet.append(["Назва", "Сума"])
    sheet.append(["Грант", 100])
    workbook.save(xlsx_path)
    workbook.close()
    xlsx_result = extract_document_file(xlsx_path)
    assert "[Sheet: Дані]" in xlsx_result.text
    assert "Грант\t100" in xlsx_result.text


def test_office_preflight_rejects_bomb_and_dtd(tmp_path: Path) -> None:
    bomb = tmp_path / "bomb.docx"
    with ZipFile(bomb, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"A" * 200_000)
    with pytest.raises(DocumentSecurityError, match="compression ratio"):
        preflight_office_archive(
            bomb,
            limits=DocumentLimits(max_compression_ratio=10),
        )

    dtd = tmp_path / "dtd.docx"
    with ZipFile(dtd, "w") as archive:
        archive.writestr(
            "word/document.xml",
            b'<!DOCTYPE x [<!ENTITY boom "x">]><x>&boom;</x>',
        )
    with pytest.raises(DocumentSecurityError, match="declarations"):
        preflight_office_archive(dtd, limits=DocumentLimits())


def test_document_source_byte_limit_precedes_parser(tmp_path: Path) -> None:
    source = tmp_path / "oversize.pdf"
    source.write_bytes(b"not-a-pdf")
    with pytest.raises(DocumentTooLargeError, match="byte limit"):
        extract_document_file(source, limits=DocumentLimits(max_source_bytes=4))


def test_raw_blob_dedup_origins_and_restart_integrity(tmp_path: Path) -> None:
    store, repository, root = _repo(tmp_path)
    first = root / "first.txt"
    second = root / "second.txt"
    first.write_text("same raw artifact", encoding="utf-8")
    second.write_text("same raw artifact", encoding="utf-8")
    blobs = ContentAddressedBlobStore(tmp_path / "blobs")
    service = LocalCorpusService(repository, allowed_root=root)

    one = service.ingest_artifact(
        SourceSpec("s1", "ws", SourceKind.LOCAL_FILE, str(first)),
        blob_store=blobs,
    )
    two = service.ingest_artifact(
        SourceSpec("s2", "ws", SourceKind.LOCAL_FILE, str(second)),
        blob_store=blobs,
    )

    assert one.artifact.artifact_id == two.artifact.artifact_id
    assert one.corpus is not None and two.corpus is not None
    assert one.corpus.document.document_id == two.corpus.document.document_id
    restarted = ResearchRepository(SQLiteStore(store.path))
    assert restarted.artifact_origin_count(one.artifact.artifact_id) == 2
    assert ContentAddressedBlobStore(blobs.root).resolve(one.artifact).is_file()


def test_blob_corruption_is_detected_on_restart(tmp_path: Path) -> None:
    _, repository, root = _repo(tmp_path)
    source = root / "source.txt"
    source.write_text("integrity", encoding="utf-8")
    blobs = ContentAddressedBlobStore(tmp_path / "blobs")
    service = LocalCorpusService(repository, allowed_root=root)
    result = service.ingest_artifact(
        SourceSpec("source", "ws", SourceKind.LOCAL_FILE, str(source)),
        blob_store=blobs,
    )
    stored = blobs.resolve(result.artifact)
    stored.write_bytes(b"x" * result.artifact.byte_size)

    with pytest.raises(BlobStoreError, match="digest"):
        ContentAddressedBlobStore(blobs.root).resolve(result.artifact)


def test_deterministic_chunking_and_persisted_chunk_order(tmp_path: Path) -> None:
    text = " ".join(f"token-{index}" for index in range(400))
    policy = ChunkPolicy(max_chars=128, overlap_chars=16)
    chunks = chunk_text(text, policy=policy)
    assert chunks == chunk_text(text, policy=policy)
    assert len(chunks) > 2
    assert all(len(chunk) <= 128 for chunk in chunks)

    _, repository, root = _repo(tmp_path)
    source = root / "long.txt"
    source.write_text("word " * 1500, encoding="utf-8")
    service = LocalCorpusService(repository, allowed_root=root)
    result = service.ingest(
        SourceSpec("long", "ws", SourceKind.LOCAL_FILE, str(source))
    )
    persisted = repository.chunk_texts(result.document.document_id)
    assert len(persisted) > 1
    assert persisted == chunk_text(result.document.text)


def test_folder_ingestion_is_ordered_bounded_and_reports_failures(tmp_path: Path) -> None:
    store, repository, root = _repo(tmp_path)
    folder = root / "folder"
    folder.mkdir()
    (folder / "b.json").write_text('{"value": "два"}', encoding="utf-8")
    (folder / "a.txt").write_text("один", encoding="utf-8")
    (folder / "broken.docx").write_bytes(b"not-a-zip")
    (folder / "ignored.bin").write_bytes(b"ignored")
    service = LocalCorpusService(repository, allowed_root=root)

    result = service.ingest_folder(
        "ws",
        folder,
        blob_store=ContentAddressedBlobStore(tmp_path / "blobs"),
    )

    titles = [
        item.corpus.document.title
        for item in result.imported
        if item.corpus is not None
    ]
    assert titles == ["a.txt", "b.json"]
    assert result.skipped_unsupported == 1
    assert len(result.failures) == 1
    assert result.failures[0].locator.endswith("broken.docx")
    assert result.failures[0].error_type == "DocumentIngestionError"
    with store.connection() as conn:
        failed = conn.execute(
            "SELECT status FROM corpus_extractions WHERE status='failed'"
        ).fetchall()
    assert len(failed) == 1
