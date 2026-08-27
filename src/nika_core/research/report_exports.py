from __future__ import annotations

import csv
import hashlib
import html
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from enum import StrEnum
from io import BytesIO, StringIO
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from openpyxl import Workbook

from nika_core.research.review import AccessibleResearchReport, ResearchCard

_FORMULA_PREFIXES = ("=", "+", "-", "@")
_OFFICE_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
_OFFICE_CREATOR = "Nika Core"


class ResearchReportFormat(StrEnum):
    TXT = "txt"
    CSV = "csv"
    HTML = "html"
    DOCX = "docx"
    XLSX = "xlsx"


@dataclass(frozen=True, slots=True)
class RenderedResearchReport:
    filename: str
    media_type: str
    content: bytes

    @property
    def sha256(self) -> str:
        return hashlib.sha256(self.content).hexdigest()


class ResearchReportExporter:
    """Render one canonical accessible research report without filesystem or network access."""

    def render(
        self,
        report: AccessibleResearchReport,
        report_format: ResearchReportFormat,
    ) -> RenderedResearchReport:
        if not isinstance(report, AccessibleResearchReport):
            raise TypeError("report must be an AccessibleResearchReport")
        if not isinstance(report_format, ResearchReportFormat):
            raise TypeError("report_format must be a ResearchReportFormat")

        filename = f"research-results-{_safe_id(report.result_set_id)}.{report_format.value}"
        if report_format is ResearchReportFormat.TXT:
            return RenderedResearchReport(
                filename, "text/plain; charset=utf-8", report.text.encode()
            )
        if report_format is ResearchReportFormat.CSV:
            return RenderedResearchReport(
                filename,
                "text/csv; charset=utf-8",
                self._csv(report).encode(),
            )
        if report_format is ResearchReportFormat.HTML:
            return RenderedResearchReport(
                filename,
                "text/html; charset=utf-8",
                self._html(report).encode(),
            )
        if report_format is ResearchReportFormat.DOCX:
            return RenderedResearchReport(
                filename,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                self._docx(report),
            )
        return RenderedResearchReport(
            filename,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            self._xlsx(report),
        )

    def _csv(self, report: AccessibleResearchReport) -> str:
        output = StringIO(newline="")
        fieldnames = _fieldnames()
        writer = csv.DictWriter(output, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()
        for row in _rows(report):
            writer.writerow({key: _spreadsheet_safe(value) for key, value in row.items()})
        return output.getvalue()

    def _html(self, report: AccessibleResearchReport) -> str:
        parts = [
            "<!doctype html>",
            '<html lang="und">',
            "<head>",
            '<meta charset="utf-8">',
            "<title>Research results</title>",
            "</head>",
            "<body>",
            "<main>",
            "<h1>Research results</h1>",
            f"<p><strong>Query:</strong> {_escape(report.query)}</p>",
            f"<p><strong>Created:</strong> {_escape(report.created_at)}</p>",
            f"<p><strong>Results:</strong> {len(report.cards)}</p>",
        ]
        for position, card in enumerate(report.cards, start=1):
            parts.extend(self._html_card(position, card))
        parts.extend(["</main>", "</body>", "</html>", ""])
        return "\n".join(parts)

    def _html_card(self, position: int, card: ResearchCard) -> list[str]:
        heading_id = f"result-{position}"
        parts = [
            f'<article aria-labelledby="{heading_id}">',
            f'<h2 id="{heading_id}">Result {position}: {_escape(card.title)}</h2>',
            "<dl>",
            f"<dt>Review</dt><dd>{_escape(card.review.state.value)}</dd>",
            f"<dt>Rank</dt><dd>{card.rank}</dd>",
            f"<dt>Why matched</dt><dd>{_escape(card.why_matched)}</dd>",
            f"<dt>Summary</dt><dd>{_escape(card.snippet)}</dd>",
        ]
        if card.review.note:
            parts.append(f"<dt>Review note</dt><dd>{_escape(card.review.note)}</dd>")
        if card.review.updated_at:
            parts.append(f"<dt>Review updated</dt><dd>{_escape(card.review.updated_at)}</dd>")
        parts.append("</dl>")
        if card.evidence:
            parts.extend(["<h3>Evidence</h3>", "<ol>"])
            for evidence in card.evidence:
                freshness = evidence.freshness.value if evidence.freshness is not None else "n/a"
                parts.extend(
                    [
                        "<li><dl>",
                        f"<dt>Source ID</dt><dd>{_escape(evidence.source_id)}</dd>",
                        f"<dt>Source kind</dt><dd>{_escape(evidence.source_kind.value)}</dd>",
                        f"<dt>Freshness</dt><dd>{_escape(freshness)}</dd>",
                        f"<dt>Location</dt><dd>{_escape(evidence.locator)}</dd>",
                        f"<dt>Observed</dt><dd>{_escape(evidence.observed_at)}</dd>",
                        "</dl></li>",
                    ]
                )
            parts.append("</ol>")
        else:
            parts.append("<p><strong>Evidence:</strong> none recorded</p>")
        parts.append("</article>")
        return parts

    def _docx(self, report: AccessibleResearchReport) -> bytes:
        document = Document()
        timestamp = _office_timestamp(report.created_at)
        properties = document.core_properties
        properties.title = "Research results"
        properties.subject = "Accessible research report"
        properties.author = _OFFICE_CREATOR
        properties.last_modified_by = _OFFICE_CREATOR
        properties.created = timestamp
        properties.modified = timestamp
        properties.revision = 1

        document.add_heading("Research results", level=1)
        _add_labeled_paragraph(document, "Query", report.query)
        _add_labeled_paragraph(document, "Created", report.created_at)
        _add_labeled_paragraph(document, "Results", str(len(report.cards)))
        for position, card in enumerate(report.cards, start=1):
            document.add_heading(f"Result {position}: {card.title}", level=2)
            _add_labeled_paragraph(document, "Review", card.review.state.value)
            _add_labeled_paragraph(document, "Rank", str(card.rank))
            _add_labeled_paragraph(document, "Why matched", card.why_matched)
            _add_labeled_paragraph(document, "Summary", card.snippet)
            if card.review.note:
                _add_labeled_paragraph(document, "Review note", card.review.note)
            if card.review.updated_at:
                _add_labeled_paragraph(document, "Review updated", card.review.updated_at)
            if card.evidence:
                document.add_heading("Evidence", level=3)
                for evidence_index, evidence in enumerate(card.evidence, start=1):
                    freshness = (
                        evidence.freshness.value if evidence.freshness is not None else "n/a"
                    )
                    paragraph = document.add_paragraph(style="List Number")
                    paragraph.add_run(f"Evidence {evidence_index}").bold = True
                    _add_labeled_paragraph(document, "Source ID", evidence.source_id)
                    _add_labeled_paragraph(document, "Source kind", evidence.source_kind.value)
                    _add_labeled_paragraph(document, "Freshness", freshness)
                    _add_labeled_paragraph(document, "Location", evidence.locator)
                    _add_labeled_paragraph(document, "Observed", evidence.observed_at)
            else:
                _add_labeled_paragraph(document, "Evidence", "none recorded")
        output = BytesIO()
        document.save(output)
        return _canonicalize_office_package(output.getvalue(), timestamp)

    def _xlsx(self, report: AccessibleResearchReport) -> bytes:
        workbook = Workbook()
        timestamp = _office_timestamp(report.created_at)
        workbook.properties.title = "Research results"
        workbook.properties.subject = "Accessible research report"
        workbook.properties.creator = _OFFICE_CREATOR
        workbook.properties.lastModifiedBy = _OFFICE_CREATOR
        workbook.properties.created = timestamp
        workbook.properties.modified = timestamp

        metadata = workbook.active
        metadata.title = "Metadata"
        metadata.append(("Field", "Value"))
        metadata.append(("Result set ID", _spreadsheet_safe(report.result_set_id)))
        metadata.append(("Workspace ID", _spreadsheet_safe(report.workspace_id)))
        metadata.append(("Query", _spreadsheet_safe(report.query)))
        metadata.append(("Created", _spreadsheet_safe(report.created_at)))
        metadata.append(("Results", len(report.cards)))
        metadata.freeze_panes = "A2"
        metadata.auto_filter.ref = metadata.dimensions

        results = workbook.create_sheet("Results")
        fieldnames = _fieldnames()
        results.append(tuple(fieldnames))
        for row in _rows(report):
            results.append(tuple(_spreadsheet_safe(row[key]) for key in fieldnames))
        results.freeze_panes = "A2"
        results.auto_filter.ref = results.dimensions

        output = BytesIO()
        workbook.save(output)
        return _canonicalize_office_package(output.getvalue(), timestamp)


def _add_labeled_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(value)


def _fieldnames() -> tuple[str, ...]:
    return (
        "result_set_id",
        "workspace_id",
        "query",
        "created_at",
        "ordinal",
        "document_id",
        "title",
        "review_state",
        "review_note",
        "review_updated_at",
        "rank",
        "why_matched",
        "summary",
        "evidence_index",
        "source_id",
        "source_kind",
        "freshness",
        "locator",
        "observed_at",
    )


def _rows(report: AccessibleResearchReport) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for card in report.cards:
        evidence_items = card.evidence or (None,)
        for evidence_index, evidence in enumerate(evidence_items, start=1):
            row: dict[str, object] = {
                "result_set_id": report.result_set_id,
                "workspace_id": report.workspace_id,
                "query": report.query,
                "created_at": report.created_at,
                "ordinal": card.ordinal,
                "document_id": card.document_id,
                "title": card.title,
                "review_state": card.review.state.value,
                "review_note": card.review.note,
                "review_updated_at": card.review.updated_at or "",
                "rank": card.rank,
                "why_matched": card.why_matched,
                "summary": card.snippet,
                "evidence_index": evidence_index if evidence is not None else "",
                "source_id": "",
                "source_kind": "",
                "freshness": "",
                "locator": "",
                "observed_at": "",
            }
            if evidence is not None:
                row.update(
                    {
                        "source_id": evidence.source_id,
                        "source_kind": evidence.source_kind.value,
                        "freshness": (
                            evidence.freshness.value if evidence.freshness is not None else "n/a"
                        ),
                        "locator": evidence.locator,
                        "observed_at": evidence.observed_at,
                    }
                )
            rows.append(row)
    return rows


def _spreadsheet_safe(value: object) -> object:
    if not isinstance(value, str):
        return value
    stripped = value.lstrip()
    if stripped.startswith(_FORMULA_PREFIXES):
        return "'" + value
    return value


def _escape(value: object) -> str:
    return html.escape(str(value), quote=True)


def _safe_id(value: str) -> str:
    cleaned = "".join(
        character if character.isalnum() or character in "-_" else "-" for character in value
    )
    cleaned = cleaned.strip("-_")
    return (cleaned or "report")[:64]


def _office_timestamp(value: str) -> datetime:
    normalized = value.strip()
    if normalized.endswith("Z"):
        normalized = normalized[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(normalized)
    except ValueError as exc:
        raise ValueError(
            "report.created_at must be an ISO-8601 timestamp for Office export"
        ) from exc
    if parsed.tzinfo is not None:
        parsed = parsed.astimezone(UTC).replace(tzinfo=None)
    return parsed.replace(microsecond=0)


def _canonicalize_office_package(content: bytes, timestamp: datetime) -> bytes:
    source_buffer = BytesIO(content)
    target_buffer = BytesIO()
    timestamp_text = timestamp.strftime("%Y-%m-%dT%H:%M:%SZ").encode()
    with ZipFile(source_buffer, "r") as source:
        members = source.infolist()
        names = [member.filename for member in members]
        if len(names) != len(set(names)):
            raise RuntimeError("Office package contains duplicate members")
        with ZipFile(target_buffer, "w", compression=ZIP_DEFLATED, compresslevel=9) as target:
            for member in sorted(members, key=lambda item: item.filename):
                payload = source.read(member.filename)
                if member.filename == "docProps/core.xml":
                    for tag in (b"created", b"modified"):
                        pattern = (
                            rb"(<dcterms:"
                            + tag
                            + rb"\b[^>]*>)[^<]*(</dcterms:"
                            + tag
                            + rb">)"
                        )
                        payload, count = re.subn(
                            pattern,
                            lambda match: match.group(1) + timestamp_text + match.group(2),
                            payload,
                            count=1,
                        )
                        if count != 1:
                            raise RuntimeError(
                                f"Office core metadata missing {tag.decode()} timestamp"
                            )
                info = ZipInfo(member.filename, date_time=_OFFICE_ZIP_TIMESTAMP)
                info.compress_type = ZIP_DEFLATED
                info.create_system = 0
                info.external_attr = 0
                target.writestr(
                    info,
                    payload,
                    compress_type=ZIP_DEFLATED,
                    compresslevel=9,
                )
    return target_buffer.getvalue()
