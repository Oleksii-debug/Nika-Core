from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime, time
from importlib.metadata import version
from pathlib import Path, PurePosixPath
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from openpyxl.xml import functions as openpyxl_xml
from pypdf import PdfReader

from nika_core.research.models import ExtractedDocument, ExtractionStatus


class DocumentIngestionError(RuntimeError):
    pass


class DocumentSecurityError(DocumentIngestionError):
    pass


class DocumentTooLargeError(DocumentIngestionError):
    pass


@dataclass(frozen=True, slots=True)
class DocumentLimits:
    max_source_bytes: int = 64 * 1024 * 1024
    max_pages: int = 500
    max_extracted_chars: int = 10_000_000
    max_zip_members: int = 2048
    max_zip_member_bytes: int = 16 * 1024 * 1024
    max_zip_total_bytes: int = 64 * 1024 * 1024
    max_compression_ratio: float = 100.0

    def __post_init__(self) -> None:
        values = (
            self.max_source_bytes,
            self.max_pages,
            self.max_extracted_chars,
            self.max_zip_members,
            self.max_zip_member_bytes,
            self.max_zip_total_bytes,
        )
        if any(value < 1 for value in values) or self.max_compression_ratio <= 1:
            raise ValueError("document limits must be positive and bounded")


_DOCUMENT_PACKAGES = {
    ".pdf": "pypdf",
    ".docx": "python-docx",
    ".xlsx": "openpyxl",
}
_DOCUMENT_MEDIA_SUFFIXES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
}


def _check_archive_member_name(name: str) -> None:
    path = PurePosixPath(name.replace("\\", "/"))
    if path.is_absolute() or ".." in path.parts:
        raise DocumentSecurityError("Office archive contains an unsafe member path")


def preflight_office_archive(path: Path | str, *, limits: DocumentLimits) -> None:
    try:
        with ZipFile(path) as archive:
            members = archive.infolist()
            if len(members) > limits.max_zip_members:
                raise DocumentTooLargeError("Office archive contains too many members")
            total = 0
            for member in members:
                _check_archive_member_name(member.filename)
                if member.flag_bits & 0x1:
                    raise DocumentSecurityError("encrypted Office archive is unsupported")
                if member.file_size > limits.max_zip_member_bytes:
                    raise DocumentTooLargeError("Office archive member exceeds size limit")
                total += member.file_size
                if total > limits.max_zip_total_bytes:
                    raise DocumentTooLargeError("Office archive exceeds expanded size limit")
                if member.file_size:
                    ratio = member.file_size / max(member.compress_size, 1)
                    if ratio > limits.max_compression_ratio:
                        raise DocumentSecurityError("Office archive compression ratio is unsafe")
                lowered = member.filename.casefold()
                if lowered.endswith((".xml", ".rels")) and member.file_size:
                    payload = archive.read(member)
                    upper = payload.upper()
                    if b"<!DOCTYPE" in upper or b"<!ENTITY" in upper:
                        raise DocumentSecurityError(
                            "Office XML declarations that can expand entities are forbidden"
                        )
    except BadZipFile as exc:
        raise DocumentIngestionError("malformed Office ZIP container") from exc


def _bounded_join(parts: list[str], *, limits: DocumentLimits) -> str:
    text = "\n".join(part for part in parts if part.strip())
    if len(text) > limits.max_extracted_chars:
        raise DocumentTooLargeError("extracted document text exceeds character limit")
    return text


def _extract_pdf(path: Path, *, limits: DocumentLimits) -> ExtractedDocument:
    try:
        reader = PdfReader(path)
        if reader.is_encrypted:
            raise DocumentSecurityError("encrypted PDF is unsupported")
        if len(reader.pages) > limits.max_pages:
            raise DocumentTooLargeError("PDF page count exceeds limit")
        parts: list[str] = []
        total = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            total += len(text)
            if total > limits.max_extracted_chars:
                raise DocumentTooLargeError("PDF extracted text exceeds character limit")
            if text.strip():
                parts.append(text)
    except (DocumentIngestionError, DocumentSecurityError, DocumentTooLargeError):
        raise
    except Exception as exc:
        raise DocumentIngestionError("PDF extraction failed") from exc

    text = "\n".join(parts)
    status = ExtractionStatus.EXTRACTED if text.strip() else ExtractionStatus.OCR_NEEDED
    return ExtractedDocument(
        title=path.name,
        text=text,
        media_type="application/pdf",
        status=status,
        extractor="pypdf",
        extractor_version=version("pypdf"),
    )


def _extract_docx(path: Path, *, limits: DocumentLimits) -> ExtractedDocument:
    preflight_office_archive(path, limits=limits)
    try:
        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [" ".join(p.text for p in cell.paragraphs).strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        text = _bounded_join(parts, limits=limits)
    except (DocumentIngestionError, DocumentSecurityError, DocumentTooLargeError):
        raise
    except Exception as exc:
        raise DocumentIngestionError("DOCX extraction failed") from exc

    return ExtractedDocument(
        title=path.name,
        text=text,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        status=ExtractionStatus.EXTRACTED if text.strip() else ExtractionStatus.EMPTY,
        extractor="python-docx",
        extractor_version=version("python-docx"),
    )


def _cell_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (datetime, date, time)):
        return value.isoformat()
    return str(value)


def _extract_xlsx(path: Path, *, limits: DocumentLimits) -> ExtractedDocument:
    preflight_office_archive(path, limits=limits)
    if not openpyxl_xml.DEFUSEDXML:
        raise DocumentSecurityError("openpyxl defusedxml protection is unavailable")
    try:
        with path.open("rb") as source:
            workbook = load_workbook(source, read_only=True, data_only=True, keep_links=False)
            try:
                parts: list[str] = []
                total = 0
                for sheet in workbook.worksheets:
                    header = f"[Sheet: {sheet.title}]"
                    parts.append(header)
                    total += len(header)
                    for row in sheet.iter_rows(values_only=True):
                        cells = [_cell_text(value).strip() for value in row]
                        if not any(cells):
                            continue
                        rendered = "\t".join(cells)
                        total += len(rendered) + 1
                        if total > limits.max_extracted_chars:
                            raise DocumentTooLargeError(
                                "XLSX extracted text exceeds character limit"
                            )
                        parts.append(rendered)
                text = "\n".join(parts)
            finally:
                workbook.close()
    except (DocumentIngestionError, DocumentSecurityError, DocumentTooLargeError):
        raise
    except Exception as exc:
        raise DocumentIngestionError("XLSX extraction failed") from exc

    return ExtractedDocument(
        title=path.name,
        text=text,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        status=ExtractionStatus.EXTRACTED if text.strip() else ExtractionStatus.EMPTY,
        extractor="openpyxl",
        extractor_version=version("openpyxl"),
    )


def _document_suffix(path: Path | str, media_type: str | None) -> str:
    if media_type is not None:
        suffix = _DOCUMENT_MEDIA_SUFFIXES.get(media_type.casefold())
        if suffix is None:
            raise DocumentIngestionError(f"unsupported document media type: {media_type}")
        return suffix
    suffix = Path(path).suffix.casefold()
    if suffix not in _DOCUMENT_PACKAGES:
        raise DocumentIngestionError(f"unsupported document format: {suffix or '<none>'}")
    return suffix


def document_extractor_identity(
    path: Path | str,
    *,
    media_type: str | None = None,
) -> tuple[str, str]:
    suffix = _document_suffix(path, media_type)
    package = _DOCUMENT_PACKAGES[suffix]
    return package, version(package)


def extract_document_file(
    path: Path | str,
    *,
    limits: DocumentLimits | None = None,
    media_type: str | None = None,
) -> ExtractedDocument:
    candidate = Path(path)
    active_limits = limits or DocumentLimits()
    if not candidate.is_file():
        raise DocumentIngestionError("document source is not a regular file")
    if candidate.stat().st_size > active_limits.max_source_bytes:
        raise DocumentTooLargeError("document source exceeds byte limit")
    suffix = _document_suffix(candidate, media_type)
    if suffix == ".pdf":
        return _extract_pdf(candidate, limits=active_limits)
    if suffix == ".docx":
        return _extract_docx(candidate, limits=active_limits)
    return _extract_xlsx(candidate, limits=active_limits)
